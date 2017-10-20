import os 
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl import load_workbook
import tkinter as tk
from tkinter import *

#district list
DList = ['BANKURA','BIRBHUM','BURDWAN','COOCHBEHAR','DAKSHIN DINAJPUR','DARJEELING','HOOGHLY','HOWRAH','JALPAIGURI','MALDA','MURSHIDABAD','NADIA','NORTH 24 PARGANAS','PASCHIM MEDINIPUR','PURBA MEDINIPUR','PURULIA','SOUTH 24 PARGANAS','UTTAR DINAJPUR']
CropList = ['AUS','AMAN','BORO','WHEAT','MAIZE','JUTE','MUSUR','MASKALAI','KHESARI','GRAM','MUSTARD','TIL','POTATO','SUGARCANE']
years = ['2013-14','2014-15','2015-16']
allCols = ['A','B','C','D','E','F','G','H','I','J','K','L','M','N','O','P','Q','R','S','T','U','V','W','X','Y','Z','AA','AB','AC','AD','AE','AF','AG','AH','AI','AJ','AK','AL','AM','AN','AO','AP','AQ','AR','AS','AT','AU','AV','AW','AX','AY','AZ']

def getList():
    lis = ent.get()
    global DList,year
    if lis != '':
        DList = []
        for wrd in (lis).split(','):
            DList.append(wrd)
    yr = enty.get()
    if yr != '':
        year=[]
        for wrd in (yr).split(','):
            year.append(wrd)
    else:
        tk.messagebox.showerror(title="Error",message="No year entered.\nPlease enter years.")
    root.destroy()
root = tk.Tk()
root.wm_title("Comparison")
lab = Label(root, width=15, text="District names", anchor='w')
ent = Entry(root, width=30)
laby = Label(root, width=15, text="Years", anchor='w')
enty = Entry(root, width=30)
inst = Label(root, justify=LEFT, fg='#656565', text="Enter district names and years separated by commas.\nEnter nothing and press okay to convert\nall districts.\nSeparate year with \'-\', for example \'2014-15\'")
but = Button(root, width=8, text="Okay", command=getList)
lab.grid(row=1,column=1)
laby.grid(row=2,column=1)
ent.grid(row=1,column=2,columnspan=2)
enty.grid(row=2,column=2,columnspan=2)
inst.grid(row=3,column=1,columnspan=2)
but.grid(row=3,column=3)
root.mainloop()

for dist in DList:

    print (dist)

    document = Document('default.docx')

    wkb=[]
    wks=[]

    style = document.styles['No Spacing']
    font = style.font
    font.size = Pt(20)
    font.bold = True
    x = document.add_paragraph()
    x.style = style
    x.add_run(dist+"\n")

    for year in years:
        
        cwb = load_workbook(filename="INPUT\\"+year+"\\"+dist.replace(" ","_")+"_Crop_"+year+".xlsx", read_only=True)
        wkb.append(cwb)
        wks.append(cwb['18.1'])

    for crop in CropList:

        print (" "+crop)

        for k in range (0,3):
            
            p = document.add_paragraph('Crop: ')
            p.add_run(crop+"\n").bold = True

            if k==0:
                p.add_run('Area').bold=True
            if k==1:
                p.add_run('Production').bold=True
            if k==2:
                p.add_run('Yield').bold=True

            table = document.add_table(rows=2, cols=5)
            table.style = 'Table Grid'

            for i in range (0,11):
                row_cells = table.add_row().cells
            
            hd00 = (table.cell(0, 0)).merge(table.cell(1, 0))
            hd01 = (table.cell(0, 1)).merge(table.cell(0, 2))
            hd02 = (table.cell(0, 3)).merge(table.cell(0, 4))
            table.cell(0,0).text = "Range of Percentage Variation"
            table.cell(0,1).text = "Name of Block"
            table.cell(0,3).text = "Number of Blocks"
            table.cell(1,1).text = years[1]+" over "+years[0]
            table.cell(1,2).text = years[2]+" over "+years[1]
            table.cell(1,3).text = years[1]+" over "+years[0]
            table.cell(1,4).text = years[2]+" over "+years[1]
            table.cell(2,0).text = "< 5"
            table.cell(3,0).text = "6 - 10"
            table.cell(4,0).text = "11 - 20"
            table.cell(5,0).text = "21 - 30"
            table.cell(6,0).text = "31 - 40"
            table.cell(7,0).text = "41 - 50"
            table.cell(8,0).text = "51 - 100"
            table.cell(9,0).text = "> 100"
            table.cell(10,0).text = "Change to 0"
            table.cell(11,0).text = "Change from 0"
            table.cell(12,0).text = "ERROR"

            for i in range (2,13):
                table.cell(i,3).text = "0"
                table.cell(i,4).text = "0"

            for i in range(0,52):
                col = allCols[i]
                if wks[0][col+str(3)].value!= None:
                    if crop in wks[0][col+str(3)].value.upper():
                        for j in range (1,32):
                            nm=None
                            for m in range (0,3):
                                if wks[m]['B'+str(5+j)].value != None:
                                    nm=wks[m]['B'+str(5+j)].value
                            if nm != None:
                                for l in range (0,2):
                                    if (wks[0+l][allCols[i+k]+str(5+j)].value=="ERROR")|(wks[1+l][allCols[i+k]+str(5+j)].value=="ERROR"):
                                        table.cell(12,l+1).text += (str(nm+", "))
                                        table.cell(12,3+l).text = str(int(table.cell(12,3+l).text)+1)

                                    elif wks[0+l][allCols[i+k]+str(5+j)].value==None:
                                        if wks[1+l][allCols[i+k]+str(5+j)].value==None:
                                            table.cell(2,l+1).text += (str(nm+", "))
                                            table.cell(2,3+l).text = str(int(table.cell(2,3+l).text)+1)
                                        elif float(wks[1+l][allCols[i+k]+str(5+j)].value)==0:
                                            table.cell(2,l+1).text += (str(nm+", "))
                                            table.cell(2,3+l).text = str(int(table.cell(2,3+l).text)+1)
                                        else:
                                            table.cell(11,l+1).text += (str(nm+", "))
                                            table.cell(11,3+l).text = str(int(table.cell(11,3+l).text)+1)
                                            table.cell(9,l+1).text += (str(nm+", "))
                                            table.cell(9,3+l).text = str(int(table.cell(9,3+l).text)+1)
                                    elif wks[1+l][allCols[i+k]+str(5+j)].value==None:
                                        if float(wks[0+l][allCols[i+k]+str(5+j)].value)==0:
                                            table.cell(2,l+1).text += (str(nm+", "))
                                            table.cell(2,3+l).text = str(int(table.cell(2,3+l).text)+1)
                                        else:
                                            table.cell(10,l+1).text += (str(nm+", "))
                                            table.cell(10,3+l).text = str(int(table.cell(10,3+l).text)+1)
                                            table.cell(9,l+1).text += (str(nm+", "))
                                            table.cell(9,3+l).text = str(int(table.cell(9,3+l).text)+1)
                                    elif float(wks[0+l][allCols[i+k]+str(5+j)].value)==0:
                                        if wks[1+l][allCols[i+k]+str(5+j)].value==0:
                                            table.cell(2,l+1).text += (str(nm+", "))
                                            table.cell(2,3+l).text = str(int(table.cell(2,3+l).text)+1)
                                        else:
                                            table.cell(11,l+1).text += (str(nm+", "))
                                            table.cell(11,3+l).text = str(int(table.cell(11,3+l).text)+1)
                                            table.cell(9,l+1).text += (str(nm+", "))
                                            table.cell(9,3+l).text = str(int(table.cell(9,3+l).text)+1)
                                    elif float(wks[1+l][allCols[i+k]+str(5+j)].value)==0:
                                        table.cell(10,l+1).text += (str(nm+", "))
                                        table.cell(10,3+l).text = str(int(table.cell(10,3+l).text)+1)
                                        table.cell(9,l+1).text += (str(nm+", "))
                                        table.cell(9,3+l).text = str(int(table.cell(9,3+l).text)+1)
                                    elif ((float(wks[0+l][allCols[i+k]+str(5+j)].value)!= 0)&(float(wks[1+l][allCols[i+k]+str(5+j)].value)!=0))&((wks[0+l][allCols[i+k]+str(5+j)].value!=None)&(wks[1+l][allCols[i+k]+str(5+j)].value!=None)):
                                        delta = abs((float(wks[1+l][allCols[i+k]+str(5+j)].value)-float(wks[0+l][allCols[i+k]+str(5+j)].value))/float(wks[0+l][allCols[i+k]+str(5+j)].value))*100
                                        if delta <= 5:
                                            table.cell(2,l+1).text += (str(nm+", "))
                                            table.cell(2,3+l).text = str(int(table.cell(2,3+l).text)+1)
                                        elif (delta > 5)&(delta<=10):
                                            table.cell(3,l+1).text += (str(nm+", "))
                                            table.cell(3,3+l).text = str(int(table.cell(3,3+l).text)+1)
                                        elif (delta > 10)&(delta<=20):
                                            table.cell(4,l+1).text += (str(nm+", "))
                                            table.cell(4,3+l).text = str(int(table.cell(4,3+l).text)+1)
                                        elif (delta > 20)&(delta<=30):
                                            table.cell(5,l+1).text += (str(nm+", "))
                                            table.cell(5,3+l).text = str(int(table.cell(5,3+l).text)+1)
                                        elif (delta > 30)&(delta<=40):
                                            table.cell(6,l+1).text += (str(nm+", "))
                                            table.cell(6,3+l).text = str(int(table.cell(6,3+l).text)+1)
                                        elif (delta > 40)&(delta<=50):
                                            table.cell(7,l+1).text += (str(nm+", "))
                                            table.cell(7,3+l).text = str(int(table.cell(7,3+l).text)+1)
                                        elif (delta > 50)&(delta<=100):
                                            table.cell(8,l+1).text += (str(nm+", "))
                                            table.cell(8,3+l).text = str(int(table.cell(8,3+l).text)+1)
                                        elif (delta > 100):
                                            table.cell(9,l+1).text += (str(nm+", "))
                                            table.cell(9,3+l).text = str(int(table.cell(9,3+l).text)+1)

            document.add_page_break()

    if not os.path.exists("OUTPUT"):
            os.makedirs("OUTPUT")

    document.save("OUTPUT\\"+dist+".docx")